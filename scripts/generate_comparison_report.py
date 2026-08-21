# -*- coding: utf-8 -*-
"""Generate Word report comparing Class Diagram + ERD vs actual implementation."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

OUTPUT = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Bao_cao_so_sanh_ClassDiagram_ERD_vs_He_thong.docx",
)


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "D9E2F3")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "BÁO CÁO SO SÁNH\n"
        "CLASS DIAGRAM – ERD\n"
        "VỚI HỆ THỐNG THỰC TẾ\n"
        "Quản lý Homestay"
    )
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"Ngày lập: {datetime.now().strftime('%d/%m/%Y')}\n")
    r2.font.size = Pt(12)
    r2 = sub.add_run("Dự án: Bao_cao_nien_luan – homestay-backend / homestay-frontend")
    r2.font.size = Pt(11)
    r2.italic = True

    doc.add_page_break()

    # 1. Mục đích
    add_heading(doc, "1. Mục đích báo cáo", 1)
    doc.add_paragraph(
        "Báo cáo này đối chiếu biểu đồ Class Diagram và thiết kế cơ sở dữ liệu (ERD) "
        "trong báo cáo niên luận với mã nguồn đang triển khai (Spring Boot JPA + MySQL). "
        "Mục tiêu: xác định phần đã khớp, phần lệch thiết kế, và đề xuất bổ sung/chỉnh sửa "
        "để tài liệu thiết kế phản ánh đúng hệ thống hiện tại."
    )

    # 2. Phạm vi
    add_heading(doc, "2. Phạm vi đối chiếu", 1)
    add_table(
        doc,
        ["Nguồn", "Mô tả"],
        [
            ["Class Diagram", "User, HostRequest, Homestay, HomestayImage, Amenity, Booking, Payment, Review + các Enum"],
            ["ERD (CSDL)", "users, host_requests, homestays, homestay_images, amenities, homestay_amenities, bookings, payments, reviews"],
            ["Hệ thống thực tế", "Entity JPA trong homestay-backend + bảng MySQL do Hibernate ddl-auto=update tạo/cập nhật"],
        ],
        [4, 12],
    )

    # 3. Tổng quan
    add_heading(doc, "3. Tổng quan kết quả", 1)
    doc.add_paragraph(
        "Kết luận nhanh: Hệ thống đã triển khai đủ khối nghiệp vụ cốt lõi theo Class Diagram và ERD. "
        "Tuy nhiên, trong quá trình phát triển (UC-03 thanh toán thủ công, UC-06/08 hoàn tiền, "
        "UC-02 bản đồ, UC-08 chính sách hoàn tiền, cron job SYSTEM), mã nguồn đã mở rộng "
        "nhiều trường và bảng mà tài liệu thiết kế gốc chưa ghi nhận. "
        "Ngược lại, một số trường có trên ERD (created_at) chưa có trong Entity Java."
    )

    add_table(
        doc,
        ["Hạng mục", "Số lượng", "Ghi chú"],
        [
            ["Entity/Bảng khớp cơ bản", "9/9", "Đủ 9 thực thể cốt lõi"],
            ["Entity/Bảng bổ sung trong code", "2", "HomestayRefundRule, host_request_images"],
            ["Trường bổ sung trong code (so với diagram gốc)", "~20+", "Chủ yếu ở Payment, Homestay, Booking"],
            ["Trường có trên ERD nhưng thiếu trong code", "3", "created_at ở users, host_requests, reviews"],
            ["Method trên Class Diagram", "0 trên Entity", "Logic nằm ở Service layer (đúng kiến trúc Spring)"],
        ],
        [5, 3, 8],
    )

    doc.add_page_break()

    # 4. Chi tiết từng entity
    add_heading(doc, "4. Đối chiếu chi tiết theo từng thực thể", 1)

    # 4.1 User
    add_heading(doc, "4.1. User / bảng users", 2)
    add_table(
        doc,
        ["Thuộc tính", "Class Diagram", "ERD", "Code hiện tại", "Đánh giá"],
        [
            ["id", "✓", "PK BIGINT", "✓ Long", "Khớp"],
            ["email", "✓ unique", "U VARCHAR", "✓ unique", "Khớp"],
            ["password", "✓", "VARCHAR", "✓ (BCrypt nên dùng)", "Khớp – cần rà soát mã hóa"],
            ["fullName", "✓", "full_name", "✓", "Khớp"],
            ["phone", "✓", "VARCHAR", "✓", "Khớp"],
            ["avatar", "✓", "VARCHAR", "✓", "Khớp"],
            ["role", "RoleEnum", "VARCHAR", "✓ RoleEnum", "Khớp (USER/HOST/ADMIN)"],
            ["bankName", "✓", "bank_name", "✓", "Khớp – phục vụ CK thủ công"],
            ["bankHolder", "✓", "bank_holder", "✓", "Khớp"],
            ["bankAccount", "✓", "bank_account", "✓", "Khớp"],
            ["created_at", "—", "✓ TIMESTAMP", "✗ chưa có", "Cần bổ sung Entity hoặc xóa khỏi ERD"],
            ["locked / khóa TK", "— (BR-1)", "—", "✗ chưa có", "API khóa tài khoản Admin chưa triển khai"],
        ],
        [3, 2, 2.5, 3, 4],
    )
    doc.add_paragraph(
        "Method trên diagram (register, login, updateProfile, submitHostRequest): "
        "được triển khai qua AuthController, UserController, HostRequestController – không đặt trên Entity User. "
        "Đây là cách làm đúng với Spring Boot; Class Diagram nên ghi chú «implemented in Service» hoặc tách lớp Service."
    )

    # 4.2 HostRequest
    add_heading(doc, "4.2. HostRequest / bảng host_requests", 2)
    add_table(
        doc,
        ["Thuộc tính", "Class Diagram", "ERD", "Code hiện tại", "Đánh giá"],
        [
            ["id", "✓", "PK", "✓", "Khớp"],
            ["userId → user", "✓ FK", "user_id FK", "✓ @ManyToOne User", "Khớp"],
            ["idCardNumber", "✓", "id_card_number", "✓", "Khớp"],
            ["licenseImageUrl", "✓", "license_image_url", "✓", "Khớp (1 ảnh cũ)"],
            ["documentImages", "—", "—", "✓ List → host_request_images", "Bổ sung thực tế – cần thêm vào ERD"],
            ["status", "RequestStatusEnum", "VARCHAR", "✓", "Khớp"],
            ["adminNote", "✓", "admin_note TEXT", "✓", "Khớp"],
            ["created_at", "—", "✓", "✗", "Cần bổ sung hoặc cập nhật ERD"],
        ],
        [3, 2, 2.5, 3, 4],
    )

    # 4.3 Homestay
    add_heading(doc, "4.3. Homestay / bảng homestays", 2)
    add_table(
        doc,
        ["Thuộc tính", "Class Diagram", "ERD", "Code hiện tại", "Đánh giá"],
        [
            ["id, hostId, title, description, address, city", "✓", "✓", "✓", "Khớp"],
            ["pricePerNight", "Decimal", "DECIMAL", "BigDecimal", "Khớp (diagram ghi String là lỗi thiết kế)"],
            ["maxGuests, status, averageRating", "✓", "✓", "✓", "Khớp"],
            ["bedrooms, beds, bathrooms", "—", "—", "✓", "Bổ sung UI – nên thêm vào diagram/ERD"],
            ["latitude, longitude", "— (spec bổ sung)", "—", "✓", "Phục vụ VietMap – nên thêm vào diagram/ERD"],
            ["refundRules", "—", "—", "✓ HomestayRefundRule", "Bảng mới – bắt buộc cập nhật ERD"],
        ],
        [3, 2, 2.5, 3, 4],
    )

    # 4.4 HomestayImage, Amenity
    add_heading(doc, "4.4. HomestayImage & Amenity", 2)
    doc.add_paragraph(
        "HomestayImage (id, homestayId, imageUrl, isPrimary) và Amenity (id, name, icon) khớp hoàn toàn "
        "giữa Class Diagram, ERD và code. Quan hệ N-N qua homestay_amenities đúng thiết kế. "
        "Lưu ý: ERD gốc ghi ameniti_id (thiếu chữ y) – nên sửa thành amenity_id cho thống nhất với JPA."
    )

    # 4.5 Booking
    add_heading(doc, "4.5. Booking / bảng bookings", 2)
    add_table(
        doc,
        ["Thuộc tính", "Class Diagram", "ERD", "Code hiện tại", "Đánh giá"],
        [
            ["id, bookingCode, guestId, homestayId", "✓", "✓", "✓", "Khớp"],
            ["checkin/checkout", "✓", "check_in/out_date", "checkinDate/checkoutDate", "Khớp (JPA snake_case)"],
            ["totalGuests, totalPrice, status", "✓", "✓", "✓", "Khớp"],
            ["note", "✓ (spec)", "—", "✓ TEXT", "Có trong code – nên bổ sung ERD"],
            ["createdAt", "—", "created_at ✓", "✓", "Code khớp ERD; Class Diagram nên bổ sung"],
        ],
        [3, 2, 2.5, 3, 4],
    )
    doc.add_paragraph(
        "BookingStatusEnum: PENDING, CONFIRM, REJECTED, CANCELLED, COMPLETED – khớp cả diagram và code."
    )

    # 4.6 Payment - most important
    add_heading(doc, "4.6. Payment / bảng payments", 2)
    doc.add_paragraph(
        "Đây là phần lệch nhiều nhất giữa thiết kế gốc và hệ thống thực tế. "
        "Diagram/ERD chỉ mô tả thanh toán cơ bản; code đã mở rộng cho thanh toán chuyển khoản thủ công "
        "và cơ chế hoàn tiền (auto refund – đang chuyển sang thủ công theo kế hoạch)."
    )
    add_table(
        doc,
        ["Thuộc tính", "Class Diagram", "ERD", "Code hiện tại", "Đánh giá"],
        [
            ["id, bookingId (1-1)", "✓", "✓ U FK", "✓ @OneToOne", "Khớp"],
            ["paymentMethod", "✓", "✓", "✓ BANK_TRANSFER", "Khớp"],
            ["transactionCode", "✓", "✓", "✓ (sinh TXN-xxx)", "Khớp – mã tham chiếu CK thủ công"],
            ["amount", "✓", "DECIMAL", "✓ Double", "Khớp"],
            ["status", "UNPAID/PAID/REFUNDED", "VARCHAR", "✓ PaymentStatusEnum", "Khớp enum"],
            ["paidAt", "✓", "✓", "✓", "Khớp"],
            ["receiverBankName/Holder/Account", "—", "—", "✓", "Snapshot STK Host – cần thêm ERD"],
            ["refundedAt", "—", "—", "✓", "Cần thêm ERD"],
            ["refundBankAccount", "—", "—", "✓", "STK khách nhận hoàn – cần thêm ERD"],
            ["refundNote, refundPercent, refundAmount", "—", "—", "✓", "UC-08 partial refund – cần thêm ERD"],
        ],
        [3, 2, 2.5, 3, 4],
    )
    doc.add_paragraph(
        "Method diagram: processPayment() → PaymentService.processPayment(); "
        "executeRefund() → PaymentService.triggerAutoRefund(). "
        "Nên cập nhật Class Diagram: thêm lớp PaymentService hoặc stereotype «Service»."
    )

    # 4.7 Review
    add_heading(doc, "4.7. Review / bảng reviews", 2)
    add_table(
        doc,
        ["Thuộc tính", "Class Diagram", "ERD", "Code", "Đánh giá"],
        [
            ["id, bookingId, guestId, homestayId", "✓", "✓", "✓", "Khớp"],
            ["rating, comment", "✓", "✓", "✓", "Khớp"],
            ["created_at", "—", "✓", "✗", "Cần bổ sung Entity"],
        ],
        [3, 2, 2.5, 3, 4],
    )

    # 4.8 HomestayRefundRule - new
    add_heading(doc, "4.8. HomestayRefundRule / homestay_refund_rules (CHƯA CÓ TRONG DIAGRAM GỐC)", 2)
    add_table(
        doc,
        ["Thuộc tính", "Class Diagram", "ERD", "Code", "Đánh giá"],
        [
            ["id", "—", "—", "✓ PK", "Bảng mới"],
            ["homestay_id", "—", "—", "✓ FK", "1 Homestay – N rules"],
            ["minHoursBefore", "—", "—", "✓ INT", "Mốc giờ trước check-in"],
            ["refundPercent", "—", "—", "✓ INT 0–100", "Phần trăm hoàn khi khách hủy"],
        ],
        [3, 2, 2.5, 3, 4],
    )

    doc.add_page_break()

    # 5. Quan hệ
    add_heading(doc, "5. Đối chiếu quan hệ (Relationships)", 1)
    add_table(
        doc,
        ["Quan hệ", "Diagram/ERD", "Code", "Khớp?"],
        [
            ["User 1 — N HostRequest", "✓", "✓", "Có"],
            ["User 1 — N Homestay (Host)", "✓", "✓", "Có"],
            ["User 1 — N Booking (Guest)", "✓", "✓", "Có"],
            ["Homestay 1 — N HomestayImage", "✓ composition", "✓ cascade", "Có"],
            ["Homestay N — N Amenity", "✓", "✓ homestay_amenities", "Có"],
            ["Homestay 1 — N Booking", "✓", "✓", "Có"],
            ["Booking 1 — 1 Payment", "✓", "✓", "Có"],
            ["Booking 1 — 0..1 Review", "✓", "✓", "Có"],
            ["Homestay 1 — N HomestayRefundRule", "—", "—", "✗ thiếu trên diagram"],
            ["HostRequest 1 — N ảnh giấy tờ", "—", "—", "✗ host_request_images"],
        ],
        [4, 4, 4, 2],
    )

    # 6. Enum
    add_heading(doc, "6. Đối chiếu Enum", 1)
    add_table(
        doc,
        ["Enum", "Diagram", "Code", "Ghi chú"],
        [
            ["RoleEnum", "USER, HOST, ADMIN", "USER, HOST, ADMIN", "Khớp"],
            ["RequestStatusEnum", "PENDING, APPROVED, REJECTED", "Giống", "Khớp"],
            ["HomestayStatusEnum", "ACTIVE, INACTIVE", "Giống", "Khớp"],
            ["BookingStatusEnum", "PENDING, CONFIRM, REJECTED, CANCELLED, COMPLETED", "Giống", "Khớp"],
            ["PaymentStatusEnum", "UNPAID, PAID, REFUNDED", "Giống", "Khớp – có thể thêm PAYMENT_PENDING, REFUND_PENDING khi chuyển thủ công 100%"],
        ],
        [3, 5, 3, 5],
    )

    # 7. Thành phần ngoài diagram
    add_heading(doc, "7. Thành phần triển khai ngoài Class Diagram gốc", 1)
    add_table(
        doc,
        ["Thành phần", "Mô tả", "Có trong diagram?", "Khuyến nghị"],
        [
            ["BookingExpiryScheduler", "Cron SYSTEM hủy đơn PENDING quá hạn", "Không", "Thêm actor SYSTEM + lớp scheduler vào diagram hoạt động"],
            ["BookingExpiryService/Processor", "Xử lý auto-cancel + refund", "Không", "Ghi trong sequence diagram UC-06"],
            ["HomestaySchedulerProperties", "Cấu hình deadline 24h/48h", "Không", "Ghi trong tài liệu cấu hình"],
            ["StatsService / Dashboard", "Thống kê doanh thu PAID", "Không", "UC-10 – bổ sung use case diagram"],
            ["PaymentService (Service layer)", "processPayment, triggerAutoRefund", "Method trên Payment entity", "Tách lớp Service trên class diagram"],
            ["Cloudinary / Upload", "Lưu ảnh Homestay, HostRequest", "Không", "Ghi chú kiến trúc triển khai"],
        ],
        [4, 5, 3, 5],
    )

    doc.add_page_break()

    # 8. Khuyến nghị
    add_heading(doc, "8. Khuyến nghị chỉnh sửa", 1)

    add_heading(doc, "8.1. Cập nhật Class Diagram (StarUML / báo cáo)", 2)
    items_cd = [
        "Thêm class HomestayRefundRule và quan hệ 1-N với Homestay.",
        "Bổ sung thuộc tính Homestay: bedrooms, beds, bathrooms, latitude, longitude.",
        "Bổ sung Booking: note, createdAt.",
        "Mở rộng Payment: receiverBankName/Holder/Account, refundedAt, refundBankAccount, refundNote, refundPercent, refundAmount.",
        "Bổ sung HostRequest.documentImages (hoặc class HostRequestImage).",
        "Bổ sung User.createdAt (nếu giữ trên ERD).",
        "Chuyển method nghiệp vụ (processPayment, executeRefund, createBooking…) sang lớp «Service» thay vì gắn trực tiếp Entity.",
        "Thêm actor SYSTEM và lớp BookingExpiryScheduler cho UC-06/BR-9.",
        "Ghi chú stereotype «manual bank transfer» cho luồng thanh toán/hoàn tiền thủ công (theo hướng triển khai mới).",
    ]
    for i, item in enumerate(items_cd, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")

    add_heading(doc, "8.2. Cập nhật ERD (CSDL)", 2)
    items_erd = [
        "Thêm bảng homestay_refund_rules (id, homestay_id FK, min_hours_before, refund_percent).",
        "Thêm bảng host_request_images (host_request_id FK, image_url) hoặc gộp vào host_requests nếu chỉ 1 ảnh.",
        "Bổ sung cột bookings.note (TEXT), đảm bảo bookings.created_at có trên diagram.",
        "Bổ sung cột payments: receiver_bank_name, receiver_bank_holder, receiver_bank_account, refunded_at, refund_bank_account, refund_note, refund_percent, refund_amount.",
        "Bổ sung cột homestays: bedrooms, beds, bathrooms, latitude, longitude.",
        "Sửa lỗi đánh máy homestay_amenities.ameniti_id → amenity_id.",
        "Đồng bộ created_at: thêm vào Entity Java (users, host_requests, reviews) HOẶC xóa khỏi ERD nếu không dùng.",
        "Khi chuyển thanh toán/hoàn tiền thủ công 100%: cân nhắc thêm payment_status PAYMENT_PENDING, REFUND_PENDING.",
    ]
    for i, item in enumerate(items_erd, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")

    add_heading(doc, "8.3. Cập nhật mã nguồn (nếu muốn khớp ERD gốc hơn)", 2)
    items_code = [
        "Thêm trường createdAt cho User, HostRequest, Review (với @CreationTimestamp hoặc set thủ công).",
        "Triển khai khóa tài khoản Admin (users.locked / enabled) theo BR-1 nếu báo cáo yêu cầu.",
        "Hoàn thiện luồng thanh toán/hoàn tiền thủ công: tách claim / confirm-received / confirm-refund.",
        "Cập nhật HOMESTAY_PROJECT_SPEC.md cho khớp diagram và ERD sau khi chỉnh.",
        "Không nên xóa các trường Payment/HomestayRefundRule đã có – đây là mở rộng hợp lý của nghiệp vụ.",
    ]
    for i, item in enumerate(items_code, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")

    # 9. Ma trận ưu tiên
    add_heading(doc, "9. Ma trận ưu tiên cập nhật tài liệu", 1)
    add_table(
        doc,
        ["Ưu tiên", "Hạng mục", "Lý do"],
        [
            ["Cao", "ERD: bảng homestay_refund_rules + cột payments mở rộng", "Lệch lớn so với CSDL thực tế"],
            ["Cao", "Class Diagram: HomestayRefundRule + Payment fields", "Báo cáo niên luận cần phản ánh UC-08"],
            ["Cao", "Diagram: actor SYSTEM + scheduler", "Đã triển khai UC-06 cron job"],
            ["Trung bình", "ERD/Diagram: Homestay bedrooms, lat/long", "Đã dùng trên FE"],
            ["Trung bình", "created_at trên users/host_requests/reviews", "ERD có, code chưa có"],
            ["Thấp", "Sửa ameniti_id → amenity_id trên ERD", "Lỗi chính tả tài liệu"],
            ["Thấp", "Tách method Entity → Service trên diagram", "Làm rõ kiến trúc Spring Boot"],
        ],
        [2, 6, 8],
    )

    # 10. Kết luận
    add_heading(doc, "10. Kết luận", 1)
    doc.add_paragraph(
        "Hệ thống thực tế KHÔNG sai so với Class Diagram/ERD gốc ở phần lõi (9 thực thể, quan hệ, enum). "
        "Sự khác biệt chủ yếu do mở rộng nghiệp vụ trong quá trình code (thanh toán CK thủ công, hoàn tiền theo %, "
        "chính sách refund theo Homestay, cron job SYSTEM, VietMap, upload nhiều ảnh HostRequest). "
        "Khuyến nghị: CẬP NHẬT Class Diagram và ERD cho khớp code (không nên cắt bớt code để khớp diagram cũ). "
        "Đồng thời bổ sung vài trường còn thiếu trong code (created_at, khóa tài khoản) nếu báo cáo yêu cầu đầy đủ ERD."
    )

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph("— Hết báo cáo —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
